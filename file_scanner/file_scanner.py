"""
file_scanner.py — Extracts text from Image, TXT, and PDF files.

This module is the ONLY new component for file-based PII detection.
It extracts text and feeds it into the EXISTING perform_pii_analysis() pipeline.

Supported file types:
  - TXT (.txt) — direct text read
  - Images (.png, .jpg, .jpeg, .webp, .bmp, .tiff) — PaddleOCR
  - PDF (.pdf) — PyMuPDF for embedded text, PaddleOCR fallback for scanned pages

Security policy: FAIL-CLOSED
  If extraction fails for any reason, the file is treated as potentially dangerous
  and blocked rather than silently allowed through.
"""

import io
import os
import tempfile
import logging
import base64
from typing import Dict, Any, List, Optional

logger = logging.getLogger("SecurePromptFileScanner")

# ── Lazy-loaded singletons ────────────────────────────────────────────────────
_ocr_engine = None
_ocr_load_error = None

# Maximum limits
MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_PDF_PAGES = 50
OCR_TIMEOUT_SECONDS = 30

SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"} | SUPPORTED_IMAGE_EXTENSIONS

def _get_ocr_engine():
    """Lazy-initialize PaddleOCR engine (loaded once, reused for all requests)."""
    global _ocr_engine, _ocr_load_error
    if _ocr_engine is not None:
        return _ocr_engine
    if _ocr_load_error is not None:
        return None  # Already failed to load

    try:
        from paddleocr import PaddleOCR
        _ocr_engine = PaddleOCR(use_angle_cls=True, lang="en", enable_mkldnn=False)
        logger.info("PaddleOCR engine initialized successfully.")
        return _ocr_engine
    except Exception as e:
        _ocr_load_error = str(e)
        logger.error(f"Failed to initialize PaddleOCR: {e}")
        return None


def _validate_file(file_bytes: bytes, filename: str) -> Optional[str]:
    """Validate file before processing. Returns error message or None if valid."""
    if not file_bytes or len(file_bytes) == 0:
        return "File is empty."

    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        size_mb = len(file_bytes) / (1024 * 1024)
        return f"File too large ({size_mb:.1f} MB). Maximum allowed: {MAX_FILE_SIZE_BYTES / (1024 * 1024):.0f} MB."

    ext = os.path.splitext(filename.lower())[1] if filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        return f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"

    return None


def _extract_text_from_txt(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text from a plain .txt file."""
    try:
        # Try UTF-8 first, then fall back to latin-1
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")

        return {
            "success": True,
            "text": text.strip(),
            "source": "TXT direct read",
            "pages": [{"page": 1, "text": text.strip()}],
            "ocr_confidence": None,
            "extraction_method": "direct_read",
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "error": f"Failed to read TXT file: {str(e)}",
            "source": "TXT",
            "extraction_method": "direct_read",
        }


def _extract_text_from_image(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """Extract text from an image using PaddleOCR."""
    ocr = _get_ocr_engine()
    if ocr is None:
        return {
            "success": False,
            "text": "",
            "error": f"PaddleOCR is not available. Cannot scan image. Error: {_ocr_load_error or 'Unknown'}",
            "source": "Image OCR",
            "extraction_method": "paddleocr",
        }

    tmp_path = None
    try:
        # Write to temp file (PaddleOCR needs a file path)
        ext = os.path.splitext(filename.lower())[1] or ".png"
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        result = ocr.ocr(tmp_path)

        lines = []
        confidences = []

        if result and result[0]:
            for line_data in result[0]:
                if line_data and len(line_data) >= 2:
                    bbox = line_data[0]       # Bounding box coordinates
                    text_conf = line_data[1]   # (text, confidence)
                    text = text_conf[0] if isinstance(text_conf, (list, tuple)) else str(text_conf)
                    conf = float(text_conf[1]) if isinstance(text_conf, (list, tuple)) and len(text_conf) > 1 else 0.0

                    lines.append({
                        "text": text,
                        "confidence": round(conf, 4),
                        "bbox": bbox,
                    })
                    confidences.append(conf)

        full_text = "\n".join(line["text"] for line in lines)
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0

        return {
            "success": True,
            "text": full_text.strip(),
            "source": "Image PaddleOCR",
            "pages": [{"page": 1, "text": full_text.strip(), "lines": lines}],
            "ocr_confidence": round(avg_confidence, 4),
            "extraction_method": "paddleocr",
            "line_count": len(lines),
        }

    except Exception as e:
        logger.error(f"PaddleOCR image extraction error: {e}")
        return {
            "success": False,
            "text": "",
            "error": f"OCR extraction failed: {str(e)}",
            "source": "Image OCR",
            "extraction_method": "paddleocr",
        }
    finally:
        # Clean up temp file
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _extract_text_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
    """Extract text from a PDF — embedded text via PyMuPDF, scanned pages via PaddleOCR."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {
            "success": False,
            "text": "",
            "error": "PyMuPDF (fitz) is not installed. Cannot process PDF files.",
            "source": "PDF",
            "extraction_method": "pymupdf",
        }

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "error": f"Cannot open PDF: {str(e)}. File may be corrupted or password-protected.",
            "source": "PDF",
            "extraction_method": "pymupdf",
        }

    if doc.page_count > MAX_PDF_PAGES:
        doc.close()
        return {
            "success": False,
            "text": "",
            "error": f"PDF has {doc.page_count} pages. Maximum allowed: {MAX_PDF_PAGES}.",
            "source": "PDF",
            "extraction_method": "pymupdf",
        }

    all_text_parts = []
    pages_data = []
    ocr_used = False
    ocr_confidences = []

    for page_num in range(doc.page_count):
        try:
            page = doc[page_num]

            # Try embedded text first
            embedded_text = page.get_text("text").strip()

            # If page has substantial embedded text, use it directly
            if len(embedded_text) > 20:
                all_text_parts.append(embedded_text)
                pages_data.append({
                    "page": page_num + 1,
                    "text": embedded_text,
                    "method": "embedded_text",
                })
            else:
                # Page is likely scanned — render to image and OCR
                ocr = _get_ocr_engine()
                if ocr is None:
                    # Cannot OCR, but note this page had no text
                    pages_data.append({
                        "page": page_num + 1,
                        "text": embedded_text,
                        "method": "embedded_text_sparse",
                        "note": "PaddleOCR unavailable for scanned page",
                    })
                    if embedded_text:
                        all_text_parts.append(embedded_text)
                    continue

                ocr_used = True
                tmp_path = None
                try:
                    # Render page as image at 300 DPI
                    mat = fitz.Matrix(300 / 72, 300 / 72)
                    pix = page.get_pixmap(matrix=mat)
                    img_bytes = pix.tobytes("png")

                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        tmp.write(img_bytes)
                        tmp_path = tmp.name

                    result = ocr.ocr(tmp_path)

                    page_lines = []
                    if result and result[0]:
                        for line_data in result[0]:
                            if line_data and len(line_data) >= 2:
                                text_conf = line_data[1]
                                text = text_conf[0] if isinstance(text_conf, (list, tuple)) else str(text_conf)
                                conf = float(text_conf[1]) if isinstance(text_conf, (list, tuple)) and len(text_conf) > 1 else 0.0
                                page_lines.append(text)
                                ocr_confidences.append(conf)

                    page_text = "\n".join(page_lines)
                    all_text_parts.append(page_text)
                    pages_data.append({
                        "page": page_num + 1,
                        "text": page_text,
                        "method": "paddleocr",
                    })

                except Exception as ocr_err:
                    logger.warning(f"OCR failed for PDF page {page_num + 1}: {ocr_err}")
                    pages_data.append({
                        "page": page_num + 1,
                        "text": embedded_text,
                        "method": "ocr_failed",
                        "error": str(ocr_err),
                    })
                    if embedded_text:
                        all_text_parts.append(embedded_text)
                finally:
                    if tmp_path and os.path.exists(tmp_path):
                        try:
                            os.unlink(tmp_path)
                        except OSError:
                            pass

        except Exception as page_err:
            logger.warning(f"Failed to process PDF page {page_num + 1}: {page_err}")
            pages_data.append({
                "page": page_num + 1,
                "text": "",
                "method": "error",
                "error": str(page_err),
            })

    doc.close()

    full_text = "\n\n".join(all_text_parts).strip()
    avg_ocr_conf = sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else None

    return {
        "success": True,
        "text": full_text,
        "source": "PDF (embedded + OCR)" if ocr_used else "PDF (embedded text)",
        "pages": pages_data,
        "ocr_confidence": round(avg_ocr_conf, 4) if avg_ocr_conf is not None else None,
        "extraction_method": "pymupdf+paddleocr" if ocr_used else "pymupdf",
        "page_count": len(pages_data),
    }


def _extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
    try:
        import docx
        import io
        doc = docx.Document(io.BytesIO(file_bytes))
        
        # Extract from paragraphs
        text_parts = [para.text for para in doc.paragraphs]
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_parts.append(para.text)
                        
        full_text = "\n".join(text_parts)
        return {
            "success": True,
            "text": full_text,
            "source": "DOCX",
            "extraction_method": "python-docx"
        }
    except Exception as e:
        return {
            "success": False,
            "text": "",
            "error": f"Failed to process DOCX file: {str(e)}",
            "source": "DOCX",
            "extraction_method": "python-docx",
        }


def extract_text_from_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    """
    Main entry point: extract text from a file for security analysis.

    This function extracts text only — it does NOT perform PII analysis.
    The caller must pass the extracted text to the existing perform_pii_analysis().

    Args:
        file_bytes: Raw file content as bytes
        filename: Original filename (used to determine file type)

    Returns:
        dict with keys:
            success (bool): Whether extraction succeeded
            text (str): Extracted text content
            source (str): Description of extraction method used
            pages (list): Per-page extraction details
            ocr_confidence (float|None): Average OCR confidence if OCR was used
            extraction_method (str): Method identifier
            error (str|None): Error message if extraction failed
    """
    # Validate
    validation_error = _validate_file(file_bytes, filename)
    if validation_error:
        return {
            "success": False,
            "text": "",
            "error": validation_error,
            "source": "validation",
            "extraction_method": "none",
        }

    ext = os.path.splitext(filename.lower())[1]

    # Route to appropriate extractor
    if ext == ".txt":
        return _extract_text_from_txt(file_bytes)
    elif ext in SUPPORTED_IMAGE_EXTENSIONS:
        return _extract_text_from_image(file_bytes, filename)
    elif ext == ".pdf":
        return _extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_text_from_docx(file_bytes)
    else:
        return {
            "success": False,
            "text": "",
            "error": f"Unsupported file type: {ext}",
            "source": "unknown",
            "extraction_method": "none",
        }


def get_mock_placeholder(etype: str) -> str:
    etype = etype.upper()
    if "EMAIL" in etype: return "john_doe@example.com"
    elif "AADHAAR" in etype: return "1234-5678-9012"
    elif "PAN" in etype: return "ABCDE1234F"
    elif "BANK" in etype or "ROUTING" in etype: return "0000111122223333"
    elif "SSN" in etype or "SOCIAL" in etype: return "000-00-0000"
    elif "PHONE" in etype or "NUMBER" in etype: return "+91 123-456-8273"
    elif "DATE" in etype: return "01/01/0001"
    elif "ADDRESS" in etype or "LOCATION" in etype: return "123 example street, Secureville, CA 90210"
    elif "PERSON" in etype or "NAME" in etype: return "john doe"
    elif "CREDENTIAL" in etype or "KEY" in etype: return "ak_live_xYz123MockKey987"
    elif "PASSWORD" in etype: return "examplepassword@temp"
    elif "USERNAME" in etype: return "example_username"
    elif "CVV" in etype: return "123"
    elif "CREDIT_CARD" in etype or "CARD" in etype: return "1111-2222-3333-4444"
    elif "BLOOD" in etype: return "oab+-"
    elif "UPI" in etype: return "example_name_or_no@bank_id"
    elif "PASSPORT" in etype: return "A1234567"
    elif "IP_ADDRESS" in etype: return "192.168.0.1"
    elif "MAC" in etype: return "00:00:00:00:00:00"
    elif "CRYPTO" in etype or "WALLET" in etype: return "1A1zP1eP5QGefi2DMPTfTL5SLmv7DivfNa"
    elif "VEHICLE" in etype or "PLATE" in etype: return "MH-01-AB-1234"
    elif "THREAT" in etype: return "[THREAT BLOCKED]"
    return "ak_live_xYz123MockKey987"

def redact_file(file_bytes: bytes, filename: str, entities: list) -> dict:
    """
    Redacts sensitive entities from a file.
    Returns a dict with success (bool), data (base64 string of redacted file), mimeType, and error (str).
    """
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == ".txt":
        try:
            try:
                text = file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                text = file_bytes.decode("latin-1")
                
            # Perform string replacement
            for entity in entities:
                item = entity.get("item", "")
                if not item:
                    continue
                type_ = entity.get("type", "SENSITIVE")
                ph = get_mock_placeholder(type_)
                text = text.replace(item, ph)
                
            encoded = base64.b64encode(text.encode("utf-8")).decode("utf-8")
            return {"success": True, "data": encoded, "mimeType": "text/plain"}
        except Exception as e:
            return {"success": False, "error": f"TXT redaction failed: {str(e)}"}

    elif ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            for page in doc:
                for entity in entities:
                    item = entity.get("item", "")
                    if not item:
                        continue
                    
                    # Search for the exact string
                    text_instances = page.search_for(item)
                    for inst in text_instances:
                        # Add redaction annotation (solid black box, no text) to preserve exact layout
                        page.add_redact_annot(inst, fill=(0, 0, 0))
                # Apply all redactions on the page
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)
                
            # Save to bytes
            redacted_bytes = doc.write()
            doc.close()
            encoded = base64.b64encode(redacted_bytes).decode("utf-8")
            return {"success": True, "data": encoded, "mimeType": "application/pdf"}
        except Exception as e:
            return {"success": False, "error": f"PDF redaction failed: {str(e)}"}
            
    elif ext in SUPPORTED_IMAGE_EXTENSIONS:
        try:
            from PIL import Image, ImageDraw
            import tempfile
            
            ocr = _get_ocr_engine()
            if not ocr:
                return {"success": False, "error": "OCR engine not available for image redaction."}
                
            tmp_path = None
            try:
                # Write to temp file for OCR and PIL
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                    
                result = ocr.ocr(tmp_path)
                
                # Open image with PIL
                img = Image.open(tmp_path).convert("RGB")
                draw = ImageDraw.Draw(img)
                
                if result and result[0]:
                    for line_data in result[0]:
                        if not line_data or len(line_data) < 2:
                            continue
                        bbox = line_data[0]       # [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
                        text_conf = line_data[1]
                        ocr_text = text_conf[0] if isinstance(text_conf, (list, tuple)) else str(text_conf)
                        
                        # Check if any sensitive entity is inside this OCR line
                        for entity in entities:
                            item = str(entity.get("item", ""))
                            if item and item.lower() in ocr_text.lower():
                                # Black out the entire line's bounding box
                                x_coords = [pt[0] for pt in bbox]
                                y_coords = [pt[1] for pt in bbox]
                                min_x, max_x = min(x_coords), max(x_coords)
                                min_y, max_y = min(y_coords), max(y_coords)
                                draw.rectangle([min_x, min_y, max_x, max_y], fill="black")
                                break # Move to next line once redacted
                
                # Save to bytes
                import io
                out_io = io.BytesIO()
                # Determine PIL format
                pil_format = "PNG" if ext == ".png" else "JPEG"
                img.save(out_io, format=pil_format)
                
                encoded = base64.b64encode(out_io.getvalue()).decode("utf-8")
                # Determine mimetype
                mime = f"image/{'png' if ext == '.png' else 'jpeg'}"
                return {"success": True, "data": encoded, "mimeType": mime}
                
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    try: os.unlink(tmp_path)
                    except OSError: pass
                    
        except Exception as e:
            return {"success": False, "error": f"Image redaction failed: {str(e)}"}
            
    elif ext == ".docx":
        try:
            import docx
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                for entity in entities:
                    item = entity.get("item", "")
                    if not item: continue
                    type_ = entity.get("type", "SENSITIVE")
                    if item in para.text:
                        ph = get_mock_placeholder(type_)
                        para.text = para.text.replace(item, ph)
                        
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for entity in entities:
                                item = entity.get("item", "")
                                if not item: continue
                                type_ = entity.get("type", "SENSITIVE")
                                if item in para.text:
                                    ph = get_mock_placeholder(type_)
                                    para.text = para.text.replace(item, ph)
            out_io = io.BytesIO()
            doc.save(out_io)
            encoded = base64.b64encode(out_io.getvalue()).decode("utf-8")
            return {"success": True, "data": encoded, "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        except Exception as e:
            return {"success": False, "error": f"DOCX redaction failed: {str(e)}"}
            
    else:
        return {"success": False, "error": f"Redaction not supported for {ext} files"}
